from lxml import etree
from docx import Document
from docx.shared import Pt
from icc.studprogs.importer.base import BaseImporter

# from docx.shared import Inches


class Importer(BaseImporter):
    def _load(self):
        self.doc = Document(self.filename)
        return self.doc

    def _as_xml(self, root, tree):
        for section in self.doc.sections:
            s = etree.SubElement(root, "section")
            s.set("type", str(section.start_type))
            # self._section(self, section, s)
        self._paragraphs(root, self.doc.paragraphs)
        for table in self.doc.tables:
            t = etree.SubElement(root, "table")
            nrows, ncols = [len(x) for x in [table.rows, table.columns]]
            t.set("num-rows", str(nrows))
            t.set("num-columns", str(ncols))
            # cell_map = {}
            for nrow in range(nrows):
                for ncol in range(ncols):
                    try:
                        cell = table.cell(nrow, ncol)
                    except IndexError:
                        # FIXME A Bug in python-focx???
                        continue
                    # c = cell_map.setdefault(ctext, etree.SubElement(t, "cell"))
                    c = etree.SubElement(t, "cell")
                    self.set(c, "x", ncol)
                    self.set(c, "y", nrow)
                    self.set(c, "w", 0)
                    self.set(c, "h", 0)
                    self.set(c, "p", -1)
                    self._paragraphs(c, cell.paragraphs)

        return self.tree

    def _paragraphs(self, parent, paragraphs):
        for par in paragraphs:
            p = etree.SubElement(parent, "par")
            par_format = par.paragraph_format
            self.set(p, "indent", par_format.first_line_indent, "0", pt=True)
            self.set(p, "left-indent", par_format.left_indent, "0", pt=True)
            self.set(p, "right-indent", par_format.right_indent, "0", pt=True)
            self.set(p, "space-before", par_format.space_before, "0", pt=True)
            self.set(p, "space-after", par_format.space_after, "0", pt=True)

            self.set(p, "keep-together", par_format.
                     keep_together)  # Paragraph have to be on the same page
            self.set(p, "keep-with-next", par_format.keep_with_next)
            self.set(p, "page-break-before", par_format.page_break_before)

            self.set(p, "alignment", par_format.alignment)
            self.set(p, "widow-control", par_format.widow_control)
            for run in par.runs:
                r = etree.SubElement(p, "style")
                if run.style is not None:
                    r.set("id", run.style.name)
                r.set("bold", "1" if run.bold else "0")
                r.set("italic", "1" if run.italic else "0")
                r.set("underline", "1" if run.underline else "0")
                self.set(r, "font-name", run.font.name)
                self.set(r, "font-size", run.font.size, pt=True)
                # TODO Color
                r.text = run.text
